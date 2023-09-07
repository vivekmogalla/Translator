from ..src.translation.google_api import translate_text
from docx import Document
from ..utils import replace_html_entities
from ..config.parameters import get_target_language


#Read the given file line by line perform translation in the target language, and write it in output file
def slm_translate_content(source_file_path, destination_file_path):
   
    try:
        # import pdb; pdb.set_trace()
        #open file in read mode
        # Load the input .docx file
        doc = Document(source_file_path)

        # Create a new Document object to store the contents
        new_doc = Document()
        
        # Iterate through paragraphs in the input document and add them to the new document
        for para in doc.paragraphs:
            #Translate the given line in target language
            Trans_input = translate_text(para.text, get_target_language())
            
            #Replace &quot; with ""
            Final_Trans_input = replace_html_entities(Trans_input)
            
            new_doc.add_paragraph(Final_Trans_input)
        
        # Save the new document to the output .docx file
        new_doc.save(destination_file_path)
        
        print("File copied successfully!")
    except FileNotFoundError:
        print("One of the files was not found.")
    except IOError:
        print("An error occurred while copying the file.")

#End of function